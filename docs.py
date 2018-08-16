# coding=utf-8
'''
=================== README =================== 
1. 安装Python3
2. 安装相应的包
        docx  --> docx需要用whl安装，python_docx-0.8.6-py2.py3-none-any.whl
        os
        logging
        time
        json
        sys

3. 配置customer.json文件，格式如下：
    客户代码：BJ_UNICOM
        CUST_NAME ( 客户中文名 )：辽宁联通
        HOSTID_AITIBASEUSER（HOSTID_用户名称）: 0x84c22609_alti1
            MEM_DB_DIR0 ( 数据文件路径 )：/altibase_dbs0
            MEM_DB_DIR1 ( 数据文件路径 )：/altibase_dbs1
            MEM_LOG_DIR0 ( 数据文件路径 )：/altibase_log
    
    例如：   
    {
    "BJ_UNICOM":{"CUST_NAME":"北京联通",
    "0xF0DF5690_altibase":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x991B40A_altibase":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"}
    },
    "LN_UNICOM":{"CUST_NAME":"辽宁联通",
    "0x84c22609_alti1":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c22609_alti2":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c22609_alti3":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c22609_alti4":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c2260a_alti1":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c2260a_alti2":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c2260a_alti3":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c2260a_alti4":{"MEM_DB_DIR0":"/altibase_dbs0","MEM_DB_DIR1":"/altibase_dbs1","MEM_LOG_DIR0":"/altibase_log"},
    "0x84c2250d_alti01":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"},
    "0x84c2250d_alti02":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"},
    "0x84c2250d_alti03":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"},
    "0x84c2250f_alti01":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"},
    "0x84c2250f_alti02":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"},
    "0x84c2250f_alti03":{"MEM_DB_DIR0":"/alti01_data01","MEM_DB_DIR1":"/alti01_data02","MEM_LOG_DIR0":"/alti01_log01"}
    }
    }
    
4. 根目录放置上一次的巡检报告及本次的巡检日志文件

5. 执行docs.py 客户代码，比如：docs.py BJ_UNICOM

6. 根目录生成新的巡检报告

7. 查看CMD窗口REP_GAP，GC_GAP 判断复制和GC情况正常与否

8. 查看altibase_sm日志，确认checkpoint情况正常与否

9. 根据结果补充巡检分析总结 
   
'''

# TODO ：目前只支持 4版本巡检日志，需要测试补充6版本巡检日志解析功能


import docx

import os
import logging
import time
import json
import sys

logger = logging.getLogger(__name__)
logger.setLevel(level = logging.INFO)
# handler = logging.FileHandler("word.log")
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# handler.setFormatter(formatter)
console = logging.StreamHandler()
# logger.addHandler(handler)
logger.addHandler(console)

O_PATH = os.path.join(os.path.abspath("."))   #本地保存目录

DTIME = time.strftime("%Y-%m-%d") 
DTIME2 = time.strftime("%Y%m%d") 

cust=""
doc=""
out_doc_name=""
MEM_DB_DIR0 = ""
MEM_DB_DIR1 = ""
MEM_LOG_DIR0 = ""

def read_docx():
    for table in doc.tables:  # 遍历所有表格
        #logger.info( '----table------')
        for row in table.rows:  # 遍历表格的所有行
            row_str = '\t'.join([cell.text for cell in row.cells])  # 一行数据
            print (row_str)
            # for cell in row.cells:
                # #logger.info( cell.text, '\t',)

# 写 docx文件
def write_docx(logfile):      
    for table in doc.tables:  # 遍历所有表格
        if len(table.rows) >2 :  # 巡检表格大于2行
            # print("HOSTID : "+logfile.HOSTID+" , ALTIBASE_USER : "+logfile.ALTIBASE_USER)
            # print("HOSTID : "+table.rows[1].cells[2].text+" , ALTIBASE_USER : "+ table.rows[1].cells[5].text )
            if table.rows[1].cells[5].text == logfile.HOSTID and table.rows[1].cells[2].text ==  logfile.ALTIBASE_USER:  # 根据HostID 识别不同的巡检表格
                for row in table.rows:  # 遍历表格的所有行                    
                    if row.cells[0].text == "数据库进程":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text 
                        # row.cells[2].text = row.cells[3].text 
                        row.cells[3].paragraphs[0].text = logfile.PID 
                        # row.cells[4].text = logfile.PID
                    elif row.cells[0].text == "开始时间":
                        row.cells[1].text = DTIME
                        row.cells[len(row.cells)-1].text = DTIME
                    elif row.cells[0].text == "CPU使用率":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text 
                        # row.cells[2].text = row.cells[3].text 
                        row.cells[3].paragraphs[0].text = logfile.CPU_USAGE 
                        # row.cells[4].text = logfile.CPU_USAGE                
                    elif row.cells[0].text.startswith("VSZ"):
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text  
                        # row.cells[2].text = row.cells[3].text 
                        row.cells[3].paragraphs[0].text  = logfile.VSZ 
                        # row.cells[4].text = logfile.VSZ                       
                    elif row.cells[0].text == "DISK(%)" and row.cells[1].text == MEM_DB_DIR0:
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.DISK_DBS0 
                    elif row.cells[0].text == "DISK(%)" and row.cells[1].text == MEM_DB_DIR1:
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text   = logfile.DISK_DBS1 
                    elif row.cells[0].text == "DISK(%)" and row.cells[1].text == MEM_LOG_DIR0:
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.DISK_LOGS                 
                    elif row.cells[0].text == "会话数量" and row.cells[1].text == "TCP":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.SESSION_TCP  
                    elif row.cells[0].text == "会话数量" and row.cells[1].text == "UNIX":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.SESSION_UNIX  
                    elif row.cells[0].text == "会话数量" and row.cells[1].text == "IPC":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.SESSION_IPC                  
                    elif row.cells[0].text.startswith("各模块内存") and row.cells[1].text == "Query":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.MEM_ALLOC_QUERY       
                    elif row.cells[0].text.startswith("各模块内存") and row.cells[1].text == "Storage":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.MEM_ALLOC_STORAGE     
                    elif row.cells[0].text.startswith("各模块内存") and row.cells[1].text == "INDEX":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.MEM_ALLOC_INDEX                     
                    elif row.cells[0].text.startswith("GC"):
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text  
                        # row.cells[2].paragraphs[0].text  = row.cells[3].paragraphs[0].text  
                        # row.cells[3].paragraphs[0].text  = logfile.GC_GAP1 
                        # row.cells[4].text = logfile.GC_GAP1   
                    elif row.cells[0].text.startswith("复制"):
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text  
                        # row.cells[2].paragraphs[0].text  = row.cells[3].paragraphs[0].text  
                        # row.cells[3].paragraphs[0].text  = logfile.REP_GAP1 
                        # row.cells[4].text = logfile.REP_GAP1    
                    elif row.cells[0].text.startswith("检查点"):
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text  
                        # row.cells[2].paragraphs[0].text  = row.cells[3].paragraphs[0].text  
                        # row.cells[3].paragraphs[0].text  = logfile.CHECKPOINT 
                        # row.cells[4].text = logfile.CHECKPOINT      
                    elif row.cells[0].text.startswith("TX") and row.cells[1].text == "Logon":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.TX_LOGON           
                    elif row.cells[0].text.startswith("TX") and row.cells[1].text == "Prepare":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.TX_PREPARE      
                    elif row.cells[0].text.startswith("TX") and row.cells[1].text == "Execute":
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[2].paragraphs[0].text   = row.cells[4].paragraphs[0].text   
                        row.cells[4].paragraphs[0].text  = logfile.TX_EXECUTE                      
                    elif row.cells[0].text.startswith("内存表空间使用量"):
                        # print(row.cells[0].text +" ===== "+row.cells[1].text)
                        row.cells[1].paragraphs[0].text = row.cells[3].paragraphs[0].text  
                        # row.cells[2].paragraphs[0].text  = row.cells[3].paragraphs[0].text    
                        row.cells[3].paragraphs[0].text  = logfile.MEM_ALLOC 
                        # row.cells[4].text = logfile.MEM_ALLOC                   
    out_doc_name=cust.get("CUST_NAME")+"定期巡检报告_"+DTIME2+".docx"
    # 另存为新的巡检报告docx文件
    doc.save(out_doc_name)
    return out_doc_name
 
 
# 巡检信息对象 
class RoundCheckInfo(object):
    def __init__(self, ALTIBASE_USER, HOSTID, HOSTNAME, PID=None, CPU_USAGE=None, VSZ=None, DISK_DBS0=None, DISK_DBS1=None, DISK_LOGS=None, SESSION_TCP=None, SESSION_UNIX=None, SESSION_IPC=None, MEM_ALLOC_QUERY=None, MEM_ALLOC_STORAGE=None, MEM_ALLOC_INDEX=None, GC_GAP1=None, GC_GAP2=None, REP_GAP1=None, REP_GAP2=None, REP_GAP3=None, CHECKPOINT=None, TX_LOGON=None, TX_PREPARE=None, TX_EXECUTE=None, MEM_ALLOC=None):
        self.ALTIBASE_USER = ALTIBASE_USER
        self.HOSTID = HOSTID
        self.HOSTNAME = HOSTNAME
        self.PID = PID
        self.CPU_USAGE = CPU_USAGE
        self.VSZ = VSZ
        self.DISK_DBS0 = DISK_DBS0
        self.DISK_DBS1 = DISK_DBS1
        self.DISK_LOGS = DISK_LOGS
        self.SESSION_TCP = SESSION_TCP
        self.SESSION_UNIX = SESSION_UNIX
        self.SESSION_IPC = SESSION_IPC
        self.MEM_ALLOC_QUERY = MEM_ALLOC_QUERY
        self.MEM_ALLOC_STORAGE = MEM_ALLOC_STORAGE
        self.MEM_ALLOC_INDEX = MEM_ALLOC_INDEX
        self.GC_GAP1 = GC_GAP1
        self.GC_GAP2 = GC_GAP2
        self.REP_GAP1 = REP_GAP1
        self.REP_GAP2 = REP_GAP2
        self.REP_GAP3 = REP_GAP3
        self.CHECKPOINT = CHECKPOINT
        self.TX_LOGON = TX_LOGON
        self.TX_PREPARE = TX_PREPARE
        self.TX_EXECUTE = TX_EXECUTE
        self.MEM_ALLOC = MEM_ALLOC
 
# 读取4版本巡检日志文件 
def  read_v4_logfile(file,cust):
    ALTIBASE_USER="" #Altibase安装用户
    HOSTID="" #HOSTID
    HOSTNAME="" #HOSTNAME
    PID=""  #数据库进程
    CPU_USAGE="" #CPU使用率
    VSZ="" #VSZ (Mb)
    DISK_DBS0="" #DISK(%) _ 1
    DISK_DBS1="" #DISK(%) _ 2
    DISK_LOGS="" #DISK(%) _ 3
    SESSION_TCP="" #会话数量 _1
    SESSION_UNIX="0" #会话数量 _2
    SESSION_IPC="0" #会话数量 _3
    MEM_ALLOC_QUERY="" #各模块内存使用量(Mb)  _1
    MEM_ALLOC_STORAGE="" #各模块内存使用量(Mb)  _2
    MEM_ALLOC_INDEX="" #各模块内存使用量(Mb)  _3
    GC_GAP1="" #GC 1
    GC_GAP2="" #GC 2
    REP_GAP1="" #复制1
    REP_GAP2="" #复制2
    REP_GAP3="" #复制3
    CHECKPOINT="" #检查点  人工确认日志
    TX_LOGON="" # TX _1
    TX_PREPARE="" # TX _2
    TX_EXECUTE="" # TX _3
    MEM_ALLOC="" #内存表空间使用量 FREE/ALLOC/MEM_MAX


    MEM_DB_DIR0 = ""
    MEM_DB_DIR1 = ""
    MEM_LOG_DIR0 = ""
    
    session_tcp_num_row_cnt=-1
    session_ipc_num_row_cnt=-1
    session_unix_num_row_cnt=-1
    gc_row_cnt=-1
    repgap_row_cnt1=-1
    repgap_row_cnt2=-1
    repgap_row_cnt3=-1
    memalloc_row_cnt=-1
    
    with open(file,'r') as f:
        for line in f.readlines():
            stmp = line.strip()
            
            if stmp.find("Host ID:") >= 0:
                stmp=stmp.replace("==>","")[len("Host ID: "):]
                HOSTID =stmp
                logger.info("HOSTID : "+HOSTID+" , ALTIBASE_USER : "+ALTIBASE_USER)      
                try :
                    MEM_DB_DIR0 = cust.get(HOSTID+"_"+ALTIBASE_USER).get("MEM_DB_DIR0")
                    MEM_DB_DIR1 = cust.get(HOSTID+"_"+ALTIBASE_USER).get("MEM_DB_DIR1")
                    MEM_LOG_DIR0 = cust.get(HOSTID+"_"+ALTIBASE_USER).get("MEM_LOG_DIR0")                
                except :
                    pass
            elif stmp.find("Altibase User: ") >= 0:
                stmp=stmp.replace("==>","")[len("Altibase User: "):]
                ALTIBASE_USER = stmp
                #logger.info("ALTIBASE_USER : "+ALTIBASE_USER)          
            elif stmp.find("HostName:") >= 0:
                stmp=stmp.replace("==>","")[len("HostName:"):]
                HOSTNAME = stmp
                #logger.info("HOSTNAME : "+HOSTNAME)       
            elif stmp.find("altibase process pid:") >= 0:
                stmp=stmp.replace("==>","")[len("altibase process pid:"):]
                PID = stmp            
                #logger.info("PID : "+PID)       
            elif stmp.find("altibase process CPU usage(%):") >= 0:
                stmp=stmp.replace("==>","")[len("altibase process CPU usage(%):"):]
                CPU_USAGE = stmp+"%"
                #logger.info("CPU_USAGE : "+CPU_USAGE)       
            elif stmp.find("altibase process VSZ:") >= 0:
                stmp=stmp.replace("==>","")[len("altibase process VSZ:"):]
                VSZ = ("%dMB" % (int(stmp)/1024))
                #logger.info("VSZ : "+VSZ)       
            elif stmp.find(MEM_DB_DIR0) >= 0 and DISK_DBS0 == "":
                stmp=stmp[-(len(MEM_DB_DIR0)+7):-len(MEM_DB_DIR0)]
                DISK_DBS0 = stmp
                # logger.info("DISK_DBS0 : "+DISK_DBS0)       
            elif stmp.find(MEM_DB_DIR1) >= 0 and DISK_DBS1 == "":
                stmp=stmp[-(len(MEM_DB_DIR1)+7):-len(MEM_DB_DIR1)]
                DISK_DBS1 = stmp
                # logger.info("DISK_DBS1 : "+DISK_DBS1)
            elif stmp.find(MEM_LOG_DIR0) >= 0 and DISK_LOGS == "":
                stmp=stmp[-(len(MEM_LOG_DIR0)+7):-len(MEM_LOG_DIR0)]
                DISK_LOGS = stmp       
                # logger.info("DISK_LOGS : "+DISK_LOGS)
            elif stmp.find("Session_tcp_sum") >= 0 :
                session_tcp_num_row_cnt = 0
                # #logger.info("Session_tcp_sum count start")
            elif session_tcp_num_row_cnt == 3:
                session_tcp_num_row_cnt=-1
                SESSION_TCP = stmp
                #logger.info("SESSION_TCP : "+SESSION_TCP)
            elif stmp.find("Session_ipc_sum") >= 0 :
                session_ipc_num_row_cnt = 0
                # #logger.info("Session_ipc_sum count start")                
            elif session_ipc_num_row_cnt == 3:
                session_ipc_num_row_cnt=-1
                SESSION_IPC = stmp
                logger.info("SESSION_IPC : "+SESSION_IPC)
            elif stmp.find("Session_unix_sum") >= 0 :
                session_unix_num_row_cnt = 0
                # #logger.info("Session_unix_sum count start")                
            elif session_unix_num_row_cnt == 3:
                session_unix_num_row_cnt=-1
                SESSION_UNIX = stmp
                logger.info("SESSION_UNIX : "+SESSION_UNIX)                
            elif stmp.find("Query_Prepare") >= 0 and MEM_ALLOC_QUERY == "":
                stmp=stmp.split()[1]
                MEM_ALLOC_QUERY = ("%dMB" % (int(stmp)/1024/1024))
                #logger.info("MEM_ALLOC_QUERY : "+MEM_ALLOC_QUERY)
            elif stmp.find("Storage_Memory_Manager") >= 0 and MEM_ALLOC_STORAGE == "":
                stmp=stmp.split()[1]
                MEM_ALLOC_STORAGE = ("%dMB" % (int(stmp)/1024/1024))
                #logger.info("MEM_ALLOC_STORAGE : "+MEM_ALLOC_STORAGE)
            elif stmp.find("Index_Memory") >= 0 and MEM_ALLOC_INDEX == "":
                stmp=stmp.split()
                if len(stmp) > 1:
                    MEM_ALLOC_INDEX = ("%dMB" % (int(stmp[1])/1024/1024))
                    #logger.info("MEM_ALLOC_INDEX : "+MEM_ALLOC_INDEX)               
            elif stmp.find("Garbage Collector Gap Exception") >= 0 :
                gc_row_cnt = 0
                # #logger.info("session_tcp_num_row_cnt count start")
            elif gc_row_cnt == 3:
                # gc_row_cnt=-1
                GC_GAP1 = stmp
                logger.info("GC_GAP1 : "+GC_GAP1)
            elif gc_row_cnt == 4:
                gc_row_cnt=-1
                GC_GAP2 = stmp
                logger.info("GC_GAP2 : "+GC_GAP2)                
            elif stmp.find("first get Replication gap value") >= 0 :
                repgap_row_cnt1 = 0
                # #logger.info("session_tcp_num_row_cnt count start")
            elif repgap_row_cnt1 == 3:
                repgap_row_cnt1=-1
                REP_GAP1 = stmp
                logger.info("REP_GAP1 : "+REP_GAP1)               
            elif stmp.find("second get Replication gap value") >= 0 :
                repgap_row_cnt2 = 0
                # #logger.info("session_tcp_num_row_cnt count start")
            elif repgap_row_cnt2 == 3:
                repgap_row_cnt2=-1
                REP_GAP2 = stmp
                logger.info("REP_GAP2 : "+REP_GAP2)      
            elif stmp.find("Third get Replication gap value") >= 0 :
                repgap_row_cnt3 = 0
                # #logger.info("session_tcp_num_row_cnt count start")
            elif repgap_row_cnt3 == 3:
                repgap_row_cnt3=-1
                REP_GAP3 = stmp
                logger.info("REP_GAP3 : "+REP_GAP3)                        
            elif stmp.find("logon current") >= 0 and TX_LOGON=="":
                stmp=stmp[len("logon current"):]
                TX_LOGON = ("%d" % int(stmp))
                #logger.info("TX_LOGON : "+TX_LOGON)                      
            elif stmp.find("execute success count") >= 0 and TX_PREPARE=="":
                stmp=stmp[len("execute success count"):]
                TX_PREPARE = ("%d" % int(stmp))
                #logger.info("TX_PREPARE : "+TX_PREPARE)      
            elif stmp.find("prepare success count") >= 0 and TX_EXECUTE=="":
                stmp=stmp[len("prepare success count"):]
                TX_EXECUTE = ("%d" % int(stmp))
                #logger.info("TX_EXECUTE : "+TX_EXECUTE)                      
            
            elif stmp.find("==>memory usage") >= 0 :
                memalloc_row_cnt = 0
                # #logger.info("session_tcp_num_row_cnt count start")
            elif memalloc_row_cnt == 3:
                memalloc_row_cnt=-1
                stmp = stmp.split()
                FREE_MEM=("%dGB" % (int(stmp[2])/1024/1024))
                ALLOC_MEM=("%dGB" % (int(stmp[1])/1024/1024))
                MAX_MEM=("%dGB" % (int(stmp[0])/1024/1024))
                MEM_ALLOC = FREE_MEM+" / " + ALLOC_MEM + " / " + MAX_MEM
                #logger.info("MEM_ALLOC : "+MEM_ALLOC)                
           
            if session_tcp_num_row_cnt >= 0 :
                session_tcp_num_row_cnt += 1
            if session_ipc_num_row_cnt >= 0 :
                session_ipc_num_row_cnt += 1
            if session_unix_num_row_cnt >= 0 :
                session_unix_num_row_cnt += 1                
            if gc_row_cnt >= 0 :
                gc_row_cnt += 1             
            if repgap_row_cnt1 >= 0 :
                repgap_row_cnt1 += 1   
            if repgap_row_cnt2 >= 0 :
                repgap_row_cnt2 += 1   
            if repgap_row_cnt3 >= 0 :
                repgap_row_cnt3 += 1                   
            if memalloc_row_cnt >= 0 :
                memalloc_row_cnt += 1          
    
    logfile = RoundCheckInfo(ALTIBASE_USER,HOSTID, HOSTNAME, PID, CPU_USAGE, VSZ, DISK_DBS0, DISK_DBS1, DISK_LOGS, SESSION_TCP, SESSION_UNIX, SESSION_IPC, MEM_ALLOC_QUERY, MEM_ALLOC_STORAGE, MEM_ALLOC_INDEX, GC_GAP1, GC_GAP2, REP_GAP1, REP_GAP2, REP_GAP3, CHECKPOINT, TX_LOGON, TX_PREPARE, TX_EXECUTE, MEM_ALLOC)
    
    out_doc_name = write_docx(logfile)
    return out_doc_name


if __name__ == "__main__":

    # 必须传入参数 客户代码比如 BJ_UNICOM
    if len(sys.argv) > 1:
        cust_code=sys.argv[1]                
    else:
        print("Please input customer code like BJ_UNICOM.")
        sys.exit(1)
    
    cust=""
    
    # 读取 customer.json ，读取 DBS_DIR0, DBS_DIR1, DBS_LOGS_DIR 路径 和 客户中文名
    if os.path.exists("./customer.json"):
        with open("./customer.json", "r", encoding='utf-8') as fj:
            try:
                customers = json.load(fj)
                if customers is not None and len(customers) > 0:
                    cust = customers.get(cust_code)
                    if cust is not None:
                        # logger.info("Customer is  %s ." % cust)
                        logger.info("Customer is  %s ." % cust.get("CUST_NAME"))
                        # 根目录放置上一次巡检报告
                        filelist = [x for x in os.listdir(O_PATH) if os.path.split(x)[1][:4]==cust.get("CUST_NAME")]
                        for file in filelist :
                            print("Org file name is : "+file)
                            # 读取上一次巡检报告，基于上一次巡检报告生成新的巡检报告
                            doc = docx.Document(os.path.join(O_PATH,file))
                            break

                    else:
                        logger.info("Customer %s is not found." % cust_code)
                        sys.exit(1)
            except Exception as e :
                print(e)
                sys.exit(1)  
    else:
        print("Please check customer.json file.")
        sys.exit(1)  

    # 根目录放置本次巡检的巡检日志文件        
    # 读取所有 巡检日志文件，后缀.log
    filelist = [x for x in os.listdir(O_PATH) if os.path.split(x)[1][-4:]=='.log']
    for file in filelist :
        out_doc_name = read_v4_logfile(file,cust)
    print("Out file name is "+out_doc_name)
